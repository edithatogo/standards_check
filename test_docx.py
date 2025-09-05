import docx
import os

# Test if we can read a DOCX file
test_file = "docx/variants/consort-cluster-crossover.docx"

if os.path.exists(test_file):
    print(f"Found test file: {test_file}")
    try:
        doc = docx.Document(test_file)
        print(f"Successfully opened document with {len(doc.paragraphs)} paragraphs")
        # Print first few paragraphs
        for i, paragraph in enumerate(doc.paragraphs[:5]):
            print(f"Paragraph {i}: {paragraph.text[:100]}")
    except Exception as e:
        print(f"Error reading document: {e}")
else:
    print(f"Test file not found: {test_file}")
    # List what files are in the directory
    docx_dir = "docx/variants"
    if os.path.exists(docx_dir):
        print(f"Files in {docx_dir}:")
        for file in os.listdir(docx_dir):
            if file.endswith(".docx"):
                print(f"  {file}")